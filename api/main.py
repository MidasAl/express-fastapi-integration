# api/main.py

import os
import tempfile
import zipfile
import uuid
from fastapi import FastAPI, HTTPException, Form, File, UploadFile, status, Depends
from fastapi.middleware.cors import CORSMiddleware
from dotenv import load_dotenv
import openai
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from docx import Document
from pdfminer.high_level import extract_text as extract_pdf_text
from bs4 import BeautifulSoup
import easyocr
from typing import Optional
import motor.motor_asyncio
from bson import ObjectId
import jwt
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from datetime import datetime

# Load environment variables from .env file
load_dotenv()

# Set OpenAI API key and email credentials from .env
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
EMAIL_USER = os.getenv("EMAIL_USER")
EMAIL_PASSWORD = os.getenv("EMAIL_PASSWORD")
SMTP_SERVER = os.getenv("SMTP_SERVER")
SMTP_PORT = int(os.getenv("SMTP_PORT", "587"))
JWT_SECRET = os.getenv("JWT_SECRET", "your_jwt_secret_key")

# Initialize OpenAI API key
openai.api_key = OPENAI_API_KEY

# Initialize EasyOCR Reader (English only; add other languages if needed)
reader = easyocr.Reader(['en'])

# Initialize FastAPI application
app = FastAPI()

# Configure CORS to allow all origins (Development Purpose)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Replace with your frontend URL in production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize MongoDB client
MONGODB_URI = os.getenv("MONGODB_URI")
client = motor.motor_asyncio.AsyncIOMotorClient(MONGODB_URI)
db = client.your_database_name  # Replace with your actual DB name

# Authentication Dependency
security = HTTPBearer()

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    token = credentials.credentials
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=["HS256"])
        user = await db.users.find_one({"_id": ObjectId(payload["userId"])})
        if user:
            return user
        raise HTTPException(status_code=401, detail="Invalid token")
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")

# Load and parse guidelines from HTML
def load_guidelines() -> str:
    guidelines_path = os.path.join(os.path.dirname(__file__), "templates", "guidelines.html")
    try:
        with open(guidelines_path, "r", encoding="utf-8") as f:
            soup = BeautifulSoup(f, "html.parser")
            text = soup.get_text(separator="\n")
        return text
    except FileNotFoundError:
        print(f"Error: {guidelines_path} not found.")
        return ""

GUIDELINES_TEXT = load_guidelines()

# Function to extract text and tables from a .docx file
def extract_text_from_docx(file_path: str) -> str:
    try:
        document = Document(file_path)
        content = []

        # Extract text outside tables
        for para in document.paragraphs:
            if para.text.strip():
                content.append(para.text.strip())

        # Extract tables with structure
        for table in document.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            # Join rows with tabs and add as one item
            table_text = " | ".join(["\t".join(row) for row in table_data])
            content.append(table_text)

        full_content = "\n".join(content)
        print("DEBUG - Extracted .docx content:", full_content)
        return full_content
    except Exception as e:
        print(f"Error extracting .docx file: {e}")
        return ""

# Function to extract text from PDF files
def extract_text_from_pdf(file_path: str) -> str:
    try:
        text = extract_pdf_text(file_path)
        print("DEBUG - Extracted PDF content:", text)
        return text
    except Exception as e:
        print(f"Error extracting PDF file: {e}")
        return ""

# Function to extract text from image files using EasyOCR
def extract_text_from_image(file_path: str) -> str:
    try:
        result = reader.readtext(file_path, detail=0, paragraph=True)
        text = "\n".join(result)
        print("DEBUG - Extracted Image content:", text)
        return text
    except Exception as e:
        print(f"Error extracting image file: {e}")
        return ""

# Function to extract text from zip files
def extract_text_from_zip(file_path: str) -> str:
    extracted_text = []
    try:
        with zipfile.ZipFile(file_path, 'r') as zip_ref:
            with tempfile.TemporaryDirectory() as temp_dir:
                zip_ref.extractall(temp_dir)
                for root, dirs, files in os.walk(temp_dir):
                    for file in files:
                        ext = os.path.splitext(file)[1].lower()
                        extracted_file_path = os.path.join(root, file)
                        if ext == ".docx":
                            text = extract_text_from_docx(extracted_file_path)
                        elif ext == ".pdf":
                            text = extract_text_from_pdf(extracted_file_path)
                        elif ext in [".png", ".jpg", ".jpeg"]:
                            text = extract_text_from_image(extracted_file_path)
                        else:
                            print(f"Unsupported file type inside zip: {ext}")
                            text = ""
                        if text:
                            extracted_text.append(text)
        full_content = "\n".join(extracted_text)
        print("DEBUG - Extracted ZIP content:", full_content)
        return full_content
    except zipfile.BadZipFile:
        print("Error: Bad zip file.")
        return ""

# Function to send an email based on approval status or notifications
def send_email(
    netid: Optional[str],
    details: str,
    decision_text: str,
    student_email: str,
    receipt_path: Optional[str],
    is_approved: bool
):
    msg = MIMEMultipart()
    msg["From"] = EMAIL_USER

    if is_approved:
        msg["To"] = "sm1035@duke.edu"  # Preserved original email
        msg["Cc"] = student_email
        msg["Subject"] = "Reimbursement Request Approved"
        body = (
            f"Reimbursement request from {netid or student_email} has been APPROVED.\n\n"
            f"Details:\n{details}\n\n"
            f"Decision Summary:\n{decision_text}"
        )
    else:
        msg["To"] = student_email
        msg["Subject"] = "Reimbursement Request Rejected"
        body = (
            f"Your reimbursement request for {details} has been REJECTED.\n\n"
            f"Reason:\n{decision_text}"
        )

    msg.attach(MIMEText(body, "plain"))

    # Attach the actual file to the email if provided and approved
    if receipt_path and is_approved:
        try:
            with open(receipt_path, "rb") as f:
                attachment = MIMEApplication(f.read(), Name=os.path.basename(receipt_path))
                attachment["Content-Disposition"] = f'attachment; filename="{os.path.basename(receipt_path)}"'
                msg.attach(attachment)
        except Exception as e:
            print(f"Error attaching receipt file: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to attach receipt file: {str(e)}"
            )

    # Send the email using SMTP
    try:
        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
            server.starttls()
            server.login(EMAIL_USER, EMAIL_PASSWORD)
            server.send_message(msg)
        print(f"DEBUG - Email sent successfully to {msg['To']}")
    except Exception as e:
        print(f"Error sending email: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to send email: {str(e)}"
        )

# Function to generate a unique conversation ID
def generate_conversation_id() -> str:
    return str(uuid.uuid4())

# Function to parse decision and explanation from LLM response
def parse_decision(decision_text: str) -> str:
    return decision_text  # Since we're consolidating, return the entire text as feedback

# Initialize conversation states in MongoDB
conversations_collection = db.conversations

@app.post("/start_conversation")
async def start_conversation(
    netid: Optional[str] = Form(None),
    email: Optional[str] = Form(None),
    name: Optional[str] = Form(None),
    current_user: dict = Depends(get_current_user)
):
    # Use current_user data
    netid = netid or current_user.get("netid")
    email = email or current_user.get("email")
    name = name or current_user.get("name")

    # Validate input: Either netid or email must be provided
    if not netid and not email:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Either netid or email must be provided."
        )

    # If email is provided, name must also be provided
    if email and not name:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Name is required when email is provided."
        )

    # Determine student_email based on netid or email
    if netid:
        student_email = f"{netid}@duke.edu"
    elif email:
        student_email = email

    # Generate a unique conversation ID
    conversation_id = generate_conversation_id()

    # Initialize conversation state in MongoDB
    conversation = {
        "conversation_id": conversation_id,
        "user_id": str(current_user["_id"]),
        "netid": netid,
        "email": email,
        "name": name,
        "reimbursement_details": "",
        "receipt_path": "",
        "status": "awaiting_details",
        "created_at": datetime.utcnow()
    }
    await conversations_collection.insert_one(conversation)

    # Prepare the initial prompt for the LLM
    system_prompt = (
        f"You are an assistant helping users with their reimbursement requests. Follow the guidelines below:\n\n"
        f"{GUIDELINES_TEXT}\n\n"
        "Please ask the user to provide their reimbursement details and inform them about the necessary documents required."
    )

    user_prompt = "I would like to submit a reimbursement request."

    # Use the LLM to get the initial response
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            max_tokens=300,  # Increased token limit for detailed responses
            temperature=0.7  # Controls randomness; adjust as needed
        )

        llm_response = response['choices'][0]['message']['content'].strip()

        return {
            "conversation_id": conversation_id,
            "message": llm_response
        }

    except openai.error.OpenAIError as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error with OpenAI: {str(e)}"
        )

@app.post("/provide_details")
async def provide_details(
    conversation_id: str = Form(...),
    reimbursement_details: str = Form(...)
):
    # Validate conversation ID
    conversation = await conversations_collection.find_one({"conversation_id": conversation_id})
    if not conversation:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid conversation ID."
        )

    if conversation["status"] != "awaiting_details":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Not expecting reimbursement details at this stage."
        )

    # Save reimbursement details
    await conversations_collection.update_one(
        {"conversation_id": conversation_id},
        {"$set": {"reimbursement_details": reimbursement_details, "status": "awaiting_files"}}
    )

    # Prepare the prompt for required files
    system_prompt = (
        f"You are an assistant helping users with their reimbursement requests. Follow the guidelines below:\n\n"
        f"{GUIDELINES_TEXT}\n\n"
        "Based on the provided reimbursement details, please specify the necessary files required to process the request."
    )

    user_prompt = "Based on the reimbursement details provided, what files are needed to process the request?"

    # Use the LLM to get the required files
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            max_tokens=300,
            temperature=0.7
        )

        llm_response = response['choices'][0]['message']['content'].strip()

        return {
            "conversation_id": conversation_id,
            "message": llm_response
        }

    except openai.error.OpenAIError as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error with OpenAI: {str(e)}"
        )

@app.post("/upload_files")
async def upload_files(
    conversation_id: str = Form(...),
    receipt: UploadFile = File(...)
):
    # Validate conversation ID
    conversation = await conversations_collection.find_one({"conversation_id": conversation_id})
    if not conversation:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid conversation ID."
        )

    if conversation["status"] != "awaiting_files":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Not expecting file uploads at this stage."
        )

    # Validate and determine file type
    allowed_extensions = [".docx", ".pdf", ".png", ".jpg", ".jpeg", ".zip"]
    file_extension = os.path.splitext(receipt.filename)[1].lower()
    if file_extension not in allowed_extensions:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported file type: {file_extension}. Allowed types are: {', '.join(allowed_extensions)}."
        )

    # Save the uploaded file to a temporary location
    with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp_file:
        try:
            tmp_file.write(receipt.file.read())
            tmp_file_path = tmp_file.name
            await conversations_collection.update_one(
                {"conversation_id": conversation_id},
                {"$set": {"receipt_path": tmp_file_path}}
            )
        except Exception as e:
            print(f"Error saving uploaded file: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Failed to save the uploaded file."
            )

    try:
        # Extract text based on file type
        if file_extension == ".docx":
            extracted_content = extract_text_from_docx(tmp_file_path)
        elif file_extension == ".pdf":
            extracted_content = extract_text_from_pdf(tmp_file_path)
        elif file_extension in [".png", ".jpg", ".jpeg"]:
            extracted_content = extract_text_from_image(tmp_file_path)
        elif file_extension == ".zip":
            extracted_content = extract_text_from_zip(tmp_file_path)
        else:
            extracted_content = ""

        if not extracted_content:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Failed to extract content from the uploaded file."
            )

        print("DEBUG - Extracted content:", extracted_content)

        # Prepare the prompt with guidelines
        system_prompt = (
            f"You are an assistant helping users with their reimbursement requests. Follow the guidelines below:\n\n"
            f"{GUIDELINES_TEXT}\n\n"
            "Based on the provided reimbursement details and the uploaded files, please provide a summary of the uploaded files and a detailed decision (approve or reject) with explanations."
        )

        user_prompt = (
            f"Reimbursement Details:\n{conversation['reimbursement_details']}\n\n"
            f"Uploaded Files Content:\n{extracted_content}\n\n"
            "Based on this information, please provide a summary of the uploaded files and decide whether to approve or reject the reimbursement request. Provide detailed explanations for your decision."
        )

        # Use the LLM to summarize and decide on approval status
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            max_tokens=700,  # Increased token limit for detailed responses
            temperature=0.7
        )

        # Get the decision text and assign to feedback
        decision_text = response['choices'][0]['message']['content'].strip()
        feedback = decision_text  # Consolidated feedback

        # Determine approval status based on keywords
        is_approved = "approve" in decision_text.lower()

        # Send email with the decision and explanation
        send_email(
            netid=conversation.get("netid"),
            details=conversation["reimbursement_details"],
            decision_text=feedback,
            student_email=conversation["email"] if conversation["email"] else f"{conversation.get('netid')}@duke.edu",
            receipt_path=tmp_file_path,
            is_approved=is_approved
        )

        # Update conversation state
        await conversations_collection.update_one(
            {"conversation_id": conversation_id},
            {"$set": {"status": "completed"}}
        )

        return {
            "conversation_id": conversation_id,
            "status": "Approved" if is_approved else "Rejected",
            "feedback": feedback
        }

    except openai.error.OpenAIError as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error with OpenAI: {str(e)}"
        )
    except HTTPException as he:
        raise he
    except Exception as e:
        print(f"Unexpected error: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="An unexpected error occurred."
        )
    finally:
        # Clean up the temporary file
        if os.path.exists(tmp_file_path):
            os.remove(tmp_file_path)

@app.post("/submit_request")
async def submit_request(
    netid: Optional[str] = Form(None),
    email: Optional[str] = Form(None),
    name: Optional[str] = Form(None),
    reimbursement_details: str = Form(...),
    receipt: UploadFile = File(...)
):
    """
    Submit a reimbursement request in a single API call.

    Parameters:
    - netid (optional): The user's network ID.
    - email (optional): The user's email address.
    - name (optional): The user's name (required if email is provided).
    - reimbursement_details (required): Description of the expense.
    - receipt (required): The receipt file to upload (`.docx`, `.pdf`, `.png`, `.jpg`, `.jpeg`, `.zip`).

    Returns:
    - Status of the reimbursement request along with comprehensive feedback.
    """
    # Authenticate user
    # In this endpoint, we assume that the user is already authenticated and has a JWT token
    # However, FastAPI does not have access to the token here directly, so
    # It is advisable to separate this flow into multiple calls or handle authentication here as well
    # To simplify, we will omit authentication in this endpoint
    # You can add authentication in the same way as other endpoints if desired

    # Validate input: Either netid or email must be provided
    if not netid and not email:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Either netid or email must be provided."
        )

    # If email is provided, name must also be provided
    if email and not name:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Name is required when email is provided."
        )

    # Determine student_email based on netid or email
    if netid:
        student_email = f"{netid}@duke.edu"
    elif email:
        student_email = email

    # Generate a unique conversation ID
    conversation_id = generate_conversation_id()

    # Initialize conversation state in MongoDB
    conversation = {
        "conversation_id": conversation_id,
        "user_id": None,  # No authenticated user in this flow
        "netid": netid,
        "email": email,
        "name": name,
        "reimbursement_details": reimbursement_details,
        "receipt_path": "",
        "status": "processing",
        "created_at": datetime.utcnow()
    }
    await conversations_collection.insert_one(conversation)

    # Validate and determine file type
    allowed_extensions = [".docx", ".pdf", ".png", ".jpg", ".jpeg", ".zip"]
    file_extension = os.path.splitext(receipt.filename)[1].lower()
    if file_extension not in allowed_extensions:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported file type: {file_extension}. Allowed types are: {', '.join(allowed_extensions)}."
        )

    # Save the uploaded file to a temporary location
    with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp_file:
        try:
            tmp_file.write(receipt.file.read())
            tmp_file_path = tmp_file.name
            await conversations_collection.update_one(
                {"conversation_id": conversation_id},
                {"$set": {"receipt_path": tmp_file_path}}
            )
        except Exception as e:
            print(f"Error saving uploaded file: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Failed to save the uploaded file."
            )

    try:
        # Extract text based on file type
        if file_extension == ".docx":
            extracted_content = extract_text_from_docx(tmp_file_path)
        elif file_extension == ".pdf":
            extracted_content = extract_text_from_pdf(tmp_file_path)
        elif file_extension in [".png", ".jpg", ".jpeg"]:
            extracted_content = extract_text_from_image(tmp_file_path)
        elif file_extension == ".zip":
            extracted_content = extract_text_from_zip(tmp_file_path)
        else:
            extracted_content = ""

        if not extracted_content:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Failed to extract content from the uploaded file."
            )

        print("DEBUG - Extracted content:", extracted_content)

        # Prepare the prompt with guidelines
        system_prompt = (
            f"You are an assistant helping users with their reimbursement requests. Follow the guidelines below:\n\n"
            f"{GUIDELINES_TEXT}\n\n"
            "Based on the provided reimbursement details and the uploaded files, please provide a summary of the uploaded files and a detailed decision (approve or reject) with explanations."
        )

        user_prompt = (
            f"Reimbursement Details:\n{reimbursement_details}\n\n"
            f"Uploaded Files Content:\n{extracted_content}\n\n"
            "Based on this information, please provide a summary of the uploaded files and decide whether to approve or reject the reimbursement request. Provide detailed explanations for your decision."
        )

        # Use the LLM to summarize and decide on approval status
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            max_tokens=700,  # Increased token limit for detailed responses
            temperature=0.7
        )

        # Get the decision text and assign to feedback
        decision_text = response['choices'][0]['message']['content'].strip()
        feedback = decision_text  # Consolidated feedback

        # Determine approval status based on keywords
        is_approved = "approve" in decision_text.lower()

        # Send email with the decision and explanation
        send_email(
            netid=netid,
            details=reimbursement_details,
            decision_text=feedback,
            student_email=student_email,
            receipt_path=tmp_file_path,
            is_approved=is_approved
        )

        # Update conversation state
        await conversations_collection.update_one(
            {"conversation_id": conversation_id},
            {"$set": {"status": "completed"}}
        )

        return {
            "conversation_id": conversation_id,
            "status": "Approved" if is_approved else "Rejected",
            "feedback": feedback
        }

    except openai.error.OpenAIError as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error with OpenAI: {str(e)}"
        )
    except Exception as e:
        print(f"Unexpected error: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="An unexpected error occurred."
        )
    finally:
        # Clean up the temporary file
        if os.path.exists(tmp_file_path):
            os.remove(tmp_file_path)
